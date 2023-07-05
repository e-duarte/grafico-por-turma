from docx import Document
import json
import pandas as pd
import os
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess

DIMESSION_4A = [7772400, 10058400]
HEADER_PATH = 'header_portugues34.json'
GRADE = 3
TITLE = 'DIAGNÓSTICO DA TURMA'

with open(HEADER_PATH, encoding="utf8") as config_file:
    header = json.load(config_file)

VARS = header['vars']
SUBJECT = header['subject']

num_vars = len(VARS)
def get_classrooms(file, grade='.'):
    classrooms_df = pd.read_csv(file, na_filter=False)
    columns = classrooms_df.columns.values

    classrooms_df = classrooms_df[classrooms_df[columns[1]] == grade] if not grade == '' else classrooms_df
    classrooms = []

    for i, row in classrooms_df.iterrows():        
        classrooms.append(row[columns[0]])    

    return classrooms

def set_orientation(doc):
    for section in doc.sections:
        new_width, new_height = DIMESSION_4A[1], DIMESSION_4A[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(.5)
        section.bottom_margin = Cm(.5)
        section.left_margin = Cm(.5)
        section.right_margin = Cm(.5)

def build_table(doc):
    table = doc.add_table(
        rows= int(num_vars) if num_vars%3 == 0 else int(num_vars/3) + 1, 
        cols=3
    )

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r in table.rows:
        for c in r.cells:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    columns = table.columns

    # libreoffice
    columns[0].width = Cm(8.5)
    columns[1].width = Cm(8.5)
    columns[2].width = Cm(8.5)

    return table

def convert_2_pdf():
    doc_list = os.listdir('docs/')
    for d in doc_list:
        output = subprocess.check_output(['libreoffice', '--convert-to', 'pdf' ,f'docs/{d}'])
        print(output)

classrooms = get_classrooms('classrooms.csv', GRADE)
path_list = os.listdir('out/')
path_list.sort()
for classroom in classrooms:
    doc = Document()
    par = doc.add_paragraph()
    run = par.add_run(f'{TITLE} {classroom}')
    run.font.bold = True
    run.font.size = Pt(18)
    par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_orientation(doc)
    set_margins(doc)
    
    table = build_table(doc)

    cells = []
    i = 0
    for r in table.rows:
        for c in r.cells:
            if i < num_vars:
                cells.append(c)
                i +=1
    for c in cells:
        img_path = path_list.pop(0)
        c.paragraphs[0].add_run().add_picture(f'out/{img_path}', width=Cm(8.5))
    
    doc.save(f'docs/{classroom}_{SUBJECT}.docx')

    convert_2_pdf()